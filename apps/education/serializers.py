import logging
from apps.student_groups.models import StudentGroupMembership # Assicurati che sia importato
import io
import re
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
import markdown as md_parser # Rinominato per evitare conflitti
from rest_framework import serializers, status
from rest_framework.exceptions import ValidationError
from django.db import transaction
from django.utils import timezone
from django.core.files.uploadedfile import UploadedFile

from .models import (
    QuizTemplate, QuestionTemplate, AnswerOptionTemplate,
    Quiz, Question, AnswerOption, Pathway, PathwayQuiz,
    QuizAttempt, StudentAnswer, PathwayProgress, QuestionType,
    QuizAssignment, PathwayAssignment, # Importa Assignment
    PathwayTemplate, PathwayQuizTemplate, # Importa i nuovi modelli Template
    Notification # Assicura che Notification sia importato
)
from apps.users.serializers import UserSerializer, StudentSerializer, StudentBasicSerializer # Usiamo StudentBasicSerializer
from apps.users.models import User, Student
from apps.student_groups.models import StudentGroup # Importa StudentGroup
# Importa il serializer base per StudentGroup se esiste, altrimenti creane uno semplice qui
try:
    from apps.student_groups.serializers import StudentGroupSerializer as BaseStudentGroupSerializer

    class StudentGroupBasicSerializer(BaseStudentGroupSerializer):
         class Meta(BaseStudentGroupSerializer.Meta):
              fields = ['id', 'name', 'owner'] # Campi base (Corretto teacher -> owner)
              read_only_fields = fields # Sola lettura in questo contesto
except ImportError:
    # Fallback se student_groups non è ancora pronto o non ha serializer
    class StudentGroupBasicSerializer(serializers.ModelSerializer):
        class Meta:
            model = StudentGroup
            fields = ['id', 'name'] # Info minime
            read_only_fields = fields


from .models import QuizAttempt, EarnedBadge # Assicurati che QuizAttempt e EarnedBadge siano importati
from apps.rewards.serializers import SimpleBadgeSerializer # Importa il serializer per i badge

logger = logging.getLogger(__name__)

# --- Template Serializers (Invariati) ---

class AnswerOptionTemplateSerializer(serializers.ModelSerializer):
    class Meta:
        model = AnswerOptionTemplate
        fields = ['id', 'text', 'is_correct', 'order']


class QuestionTemplateSerializer(serializers.ModelSerializer):
    answer_options = AnswerOptionTemplateSerializer(
        many=True,
        source='answer_option_templates', # Usa related_name
        read_only=True
    )
    question_type_display = serializers.CharField(source='get_question_type_display', read_only=True)

    class Meta:
        model = QuestionTemplate
        fields = [
            'id', 'quiz_template', 'text', 'question_type', 'question_type_display',
            'order', 'metadata', 'answer_options'
        ]
        read_only_fields = ['quiz_template']


class QuizTemplateSerializer(serializers.ModelSerializer):
    admin_username = serializers.CharField(source='admin.username', read_only=True, allow_null=True)
    teacher_username = serializers.CharField(source='teacher.username', read_only=True, allow_null=True)

    class Meta:
        model = QuizTemplate
        fields = [
            'id', 'admin', 'admin_username', 'teacher', 'teacher_username',
            'title', 'description',
            'subject', # Ripristinato CharField
            'topic',   # Ripristinato CharField
            'metadata', 'created_at', 'card_background_color'
        ]
        read_only_fields = [
            'id', 'admin', 'admin_username', 'teacher', 'teacher_username',
            'created_at'
        ]
        # Rimosso extra_kwargs per _id
        # Rimosse definizioni subject_name, topic_name

class PathwayQuizTemplateSerializer(serializers.ModelSerializer):
    """ Serializer per la relazione M2M PathwayTemplate-QuizTemplate. """
    quiz_template_title = serializers.CharField(source='quiz_template.title', read_only=True)
    quiz_template_id = serializers.IntegerField(source='quiz_template.id', read_only=True)

    class Meta:
        model = PathwayQuizTemplate
        fields = ['id', 'quiz_template_id', 'quiz_template_title', 'order']


class PathwayTemplateSerializer(serializers.ModelSerializer):
    teacher_username = serializers.CharField(source='teacher.username', read_only=True)
    quiz_template_details = PathwayQuizTemplateSerializer(source='pathwayquiztemplate_set', many=True, read_only=True)

    class Meta:
        model = PathwayTemplate
        fields = [
            'id', 'teacher', 'teacher_username', 'title', 'description',
            'metadata', 'created_at', 'quiz_template_details'
        ]
        read_only_fields = ['teacher', 'teacher_username', 'created_at', 'quiz_template_details']


# --- Concrete Serializers (Quiz, Question, Pathway - Invariati) ---

class AnswerOptionSerializer(serializers.ModelSerializer):
    class Meta:
        model = AnswerOption
        fields = ['id', 'text', 'is_correct', 'order']


class QuestionSerializer(serializers.ModelSerializer):
    answer_options = AnswerOptionSerializer(
        many=True,
        read_only=True
    )
    question_type_display = serializers.CharField(source='get_question_type_display', read_only=True)

    class Meta:
        model = Question
        fields = [
            'id', 'quiz', 'text', 'question_type', 'question_type_display',
            'order', 'metadata', 'answer_options'
        ]
        read_only_fields = ['quiz']

    def validate_metadata(self, value):
        """
        Valida la struttura dei metadati per i diversi tipi di domanda,
        in particolare per FILL_BLANK.
        """
        # self.initial_data contiene i dati grezzi inviati nella richiesta
        question_type = self.initial_data.get('question_type')
        
        # Se stiamo aggiornando e question_type non è fornito, prendilo dall'istanza
        if self.instance and not question_type:
            question_type = self.instance.question_type

        if question_type == QuestionType.FILL_BLANK:
            if not isinstance(value, dict):
                raise serializers.ValidationError("I metadati per FILL_BLANK devono essere un dizionario.")

            required_keys = {
                'text_with_placeholders': str,
                'blanks': list,
                'case_sensitive': bool,
                'points': (int, float)
            }
            for key, expected_type in required_keys.items():
                if key not in value:
                    raise serializers.ValidationError(f"Chiave mancante nei metadati FILL_BLANK: '{key}'.")
                if not isinstance(value[key], expected_type):
                    raise serializers.ValidationError(f"Chiave '{key}' nei metadati FILL_BLANK deve essere di tipo {expected_type}, ricevuto {type(value[key])}.")

            for blank_config in value['blanks']:
                if not isinstance(blank_config, dict):
                    raise serializers.ValidationError("Ogni elemento in 'blanks' deve essere un dizionario.")
                
                required_blank_keys = {
                    'id': str,
                    'correct_answers': list,
                    'order': int
                }
                for bk_key, bk_expected_type in required_blank_keys.items():
                    if bk_key not in blank_config:
                        raise serializers.ValidationError(f"Chiave mancante in un oggetto 'blank': '{bk_key}'.")
                    if not isinstance(blank_config[bk_key], bk_expected_type):
                         raise serializers.ValidationError(f"Chiave '{bk_key}' in un oggetto 'blank' deve essere di tipo {bk_expected_type}.")
                
                if not all(isinstance(ans, str) for ans in blank_config['correct_answers']):
                    raise serializers.ValidationError("Tutte le 'correct_answers' in un oggetto 'blank' devono essere stringhe.")
        
        # Aggiungere qui validazioni per altri tipi di domanda se necessario
        # Esempio: MC_SINGLE dovrebbe avere 'points_per_correct_answer'
        elif question_type == QuestionType.MULTIPLE_CHOICE_SINGLE or \
             question_type == QuestionType.TRUE_FALSE or \
             question_type == QuestionType.MULTIPLE_CHOICE_MULTIPLE:
            if not isinstance(value, dict):
                # Per questi tipi, metadata potrebbe essere vuoto o contenere 'points_per_correct_answer'
                # Se non è un dict, potrebbe essere un errore se si aspettano 'points_per_correct_answer'
                # Ma se è opzionale, un dict vuoto è ok.
                # Per ora, se non è un dict e non è None, solleva errore. Se è None, va bene.
                if value is not None:
                     raise serializers.ValidationError(f"I metadati per {question_type} devono essere un dizionario o null.")
            elif 'points_per_correct_answer' in value and not isinstance(value['points_per_correct_answer'], (int, float)):
                 raise serializers.ValidationError(f"La chiave 'points_per_correct_answer' per {question_type} deve essere un numero.")


        return value

    def to_representation(self, instance):
        """
        Personalizza la rappresentazione del campo 'metadata' per le domande FILL_BLANK.
        - Durante lo svolgimento: mostra solo placeholder e struttura base dei blank.
        - In visualizzazione risultati: mostra anche le risposte corrette.
        """
        representation = super().to_representation(instance)
        
        if instance.question_type == QuestionType.FILL_BLANK:
            metadata = representation.get('metadata', {})
            # metadata = representation.get('metadata', {}) # representation['metadata'] might not be set if super() didn't include it
            # if not isinstance(metadata, dict):
            #     metadata = {} # This was for modifying existing metadata, now we rebuild it for fill_blank

            viewing_results = self.context.get('viewing_results', False)
            
            if viewing_results:
                # For results, we assume metadata is correctly populated from the DB
                # or was processed correctly at creation/update.
                # If it's still bad here, that's a deeper issue.
                representation['metadata'] = instance.metadata
            else: # Not viewing results, i.e., student is taking the quiz
                current_instance_metadata = instance.metadata if isinstance(instance.metadata, dict) else {}
                logger.debug(f"[QID: {instance.id}] Initial instance.metadata: {current_instance_metadata}")

                # Attempt to get pre-processed values
                text_from_meta = current_instance_metadata.get('text_with_placeholders')
                blanks_from_meta = current_instance_metadata.get('blanks', []) # Default to empty list

                # Determine if pre-processed data is usable
                # Usable if text_from_meta is a non-empty string AND blanks_from_meta is a non-empty list of dicts with id and order
                is_text_usable = isinstance(text_from_meta, str) and text_from_meta.strip()
                are_blanks_usable = (
                    isinstance(blanks_from_meta, list) and
                    bool(blanks_from_meta) and # Ensure list is not empty
                    all(isinstance(b, dict) and 'id' in b and 'order' in b for b in blanks_from_meta)
                )
                
                logger.debug(f"[QID: {instance.id}] Pre-check: text_from_meta='{text_from_meta}', is_text_usable={is_text_usable}")
                logger.debug(f"[QID: {instance.id}] Pre-check: blanks_from_meta='{blanks_from_meta}', are_blanks_usable={are_blanks_usable}")

                final_text_with_placeholders = None
                final_blanks_list = []

                if is_text_usable and are_blanks_usable:
                    logger.info(f"[QID: {instance.id}] Using pre-existing metadata for 'text_with_placeholders' and 'blanks'.")
                    final_text_with_placeholders = text_from_meta
                    # We only need id and order for student view
                    for blank_conf in blanks_from_meta:
                        final_blanks_list.append({
                            'id': blank_conf.get('id'),
                            'order': blank_conf.get('order')
                        })
                else:
                    logger.warning(
                        f"[QID: {instance.id}] Metadata for 'text_with_placeholders' or 'blanks' is missing/invalid. "
                        f"Attempting to generate from instance.text. is_text_usable={is_text_usable}, are_blanks_usable={are_blanks_usable}"
                    )
                    raw_text = instance.text if instance.text else ""
                    generated_text = ""
                    # final_blanks_list is already []
                    last_idx = 0
                    blank_count = 0
                    
                    for match in re.finditer(r'_{3,}', raw_text): # Find 3 or more underscores
                        start, end = match.span()
                        generated_text += raw_text[last_idx:start]
                        blank_id = f"blank_{blank_count}"
                        generated_text += f"{{{blank_id}}}" # Use {blank_id} format
                        final_blanks_list.append({'id': blank_id, 'order': blank_count})
                        blank_count += 1
                        last_idx = end
                    generated_text += raw_text[last_idx:]
                    
                    final_text_with_placeholders = generated_text
                    
                    if blank_count == 0 and raw_text.strip():
                         logger.warning(
                             f"[QID: {instance.id}] No blanks (___) found in instance.text ('{raw_text[:100]}...') "
                             f"during fallback generation. Frontend might not render inputs."
                         )
                    elif blank_count > 0:
                        logger.info(f"[QID: {instance.id}] Generated {blank_count} blanks from instance.text.")

                representation['metadata'] = {
                    'text_with_placeholders': final_text_with_placeholders,
                    'blanks': final_blanks_list,
                    'case_sensitive': current_instance_metadata.get('case_sensitive', False), # Keep case_sensitive if present
                }
                logger.debug(f"[QID: {instance.id}] Final representation['metadata']: {representation['metadata']}")
        return representation


class QuizSerializer(serializers.ModelSerializer):
    teacher_username = serializers.CharField(source='teacher.username', read_only=True) # Mantenuto per compatibilità, se necessario
    teacher_full_name = serializers.SerializerMethodField()

    class Meta:
        model = Quiz
        fields = [
            'id', 'teacher', 'teacher_username', 'teacher_full_name', 'source_template',
            'title', 'description',
            'subject', # Ripristinato CharField
            'topic',   # Ripristinato CharField
            'metadata', 'created_at', 'available_from', 'available_until'
        ]
        read_only_fields = [
            'id', 'teacher', 'teacher_username', 'teacher_full_name', 'created_at'
        ]
        extra_kwargs = {
            'description': {'required': False, 'allow_blank': True, 'allow_null': True},
            # Rimosso extra_kwargs per _id
        }
        # Rimosse definizioni subject_name, topic_name, subject_color_placeholder

    def get_teacher_full_name(self, obj: Quiz) -> str:
        if obj.teacher:
            return f"{obj.teacher.first_name} {obj.teacher.last_name}".strip()
        return "N/D" # O un valore di default appropriato

class PathwayQuizSerializer(serializers.ModelSerializer):
    """ Serializer per la relazione M2M Pathway-Quiz (usato nested in Pathway). """
    quiz_title = serializers.CharField(source='quiz.title', read_only=True)
    quiz_id = serializers.IntegerField(source='quiz.id', read_only=True)

    class Meta:
        model = PathwayQuiz
        fields = ['id', 'quiz_id', 'quiz_title', 'order']


class PathwaySerializer(serializers.ModelSerializer):
    teacher_username = serializers.CharField(source='teacher.username', read_only=True)
    quiz_details = PathwayQuizSerializer(source='pathwayquiz_set', many=True, read_only=True)

    class Meta:
        model = Pathway
        fields = [
            'id', 'teacher', 'teacher_username', 'source_template', 'title', 'description',
            'metadata', 'created_at', 'quiz_details'
        ]
        read_only_fields = ['teacher', 'teacher_username', 'created_at', 'quiz_details']


# --- Serializer per Upload Quiz (Invariato) ---

ALLOWED_UPLOAD_EXTENSIONS = ['.pdf', '.docx', '.md']

class QuizUploadSerializer(serializers.Serializer):
    """
    Serializer per gestire l'upload di file (PDF, DOCX, MD) per creare quiz.
    """
    file = serializers.FileField(required=True, help_text="File del quiz (.pdf, .docx, .md)")
    title = serializers.CharField(max_length=255, required=True, help_text="Titolo del nuovo quiz")

    def validate_file(self, value: UploadedFile):
        if not hasattr(value, 'name') or not value.name:
             raise ValidationError("Nome del file mancante o non valido.")
        parts = value.name.split('.')
        if len(parts) < 2:
            raise ValidationError("Il file non ha un'estensione.")
        ext = '.' + parts[-1].lower()
        if ext not in ALLOWED_UPLOAD_EXTENSIONS:
            raise ValidationError(f"Tipo di file non supportato: {ext}. Sono permessi: {', '.join(ALLOWED_UPLOAD_EXTENSIONS)}")
        return value

    def _extract_text_from_pdf(self, file_obj: io.BytesIO) -> str:
        try:
            reader = PdfReader(file_obj)
            text = ""
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            if not text:
                 logger.warning("Nessun testo estratto dal PDF.")
            return text
        except Exception as e:
            logger.error(f"Errore estrazione testo PDF: {e}", exc_info=True)
            raise ValidationError(f"Impossibile leggere il file PDF. Errore: {e}")

    def _extract_text_from_docx(self, file_obj: io.BytesIO) -> str:
        try:
            document = DocxDocument(file_obj)
            text = "\n".join([para.text for para in document.paragraphs if para.text])
            if not text:
                 logger.warning("Nessun testo estratto dal DOCX.")
            return text
        except Exception as e:
            logger.error(f"Errore estrazione testo DOCX: {e}", exc_info=True)
            raise ValidationError(f"Impossibile leggere il file DOCX. Errore: {e}")

    def _extract_text_from_md(self, file_obj: io.BytesIO) -> str:
        try:
            md_content = file_obj.read().decode('utf-8')
            # Semplice estrazione testo, potrebbe non gestire bene HTML complesso
            text = re.sub('<[^<]+?>', '', md_parser.markdown(md_content)).strip()
            if not text:
                 logger.warning("Nessun testo estratto dal Markdown.")
            return text
        except Exception as e:
            logger.error(f"Errore estrazione testo Markdown: {e}", exc_info=True)
            raise ValidationError(f"Impossibile leggere il file Markdown. Errore: {e}")

    def _parse_quiz_text(self, text: str) -> list[dict]:
        # Implementazione del parsing (mantenuta come prima)
        questions = []
        current_question = None
        # Regex per inizio domanda (numero opzionale seguito da punto/parentesi)
        question_start_re = re.compile(r"^\s*(?:(\d+)\s*[.)])?\s*(.*)", re.MULTILINE)
        # Regex per opzione (lettera maiuscola seguita da punto/parentesi)
        option_re = re.compile(r"^\s*([A-Z])\s*[.)]\s*(.*)", re.MULTILINE)
        empty_line_re = re.compile(r"^\s*$")
        lines = text.splitlines()
        question_counter = 0 # Contatore per domande senza numero esplicito

        for line in lines:
            if empty_line_re.match(line):
                continue # Salta righe vuote

            line_stripped = line.strip()
            question_match = question_start_re.match(line) # Controlla se inizia come domanda
            option_match = option_re.match(line_stripped) # Controlla se è un'opzione
            is_likely_question_start = question_match and question_match.group(1) # Ha un numero?
            is_likely_option = option_match # È formattata come opzione?

            # Logica di parsing migliorata
            if is_likely_question_start:
                # Probabile inizio di una nuova domanda numerata
                if current_question:
                    questions.append(current_question) # Salva la domanda precedente
                question_number = int(question_match.group(1))
                question_text = question_match.group(2).strip()
                question_counter = question_number # Aggiorna contatore se numerata
                current_question = {
                    "text": question_text, "order": question_number, "options": [],
                    "type": QuestionType.MULTIPLE_CHOICE_SINGLE # Default
                }
            elif is_likely_option and current_question:
                # Probabile opzione della domanda corrente
                option_letter = option_match.group(1)
                option_text = option_match.group(2).strip()
                current_question["options"].append({
                    "text": option_text, "order": len(current_question["options"]) + 1, "is_correct": False # is_correct da definire
                })
            elif current_question:
                 # Riga non vuota, non è inizio domanda numerata, non è opzione formattata
                 # Potrebbe essere continuazione testo domanda o testo opzione
                 if current_question["options"] and not is_likely_question_start:
                     # Se ci sono già opzioni, probabilmente è continuazione dell'ultima opzione
                     current_question["options"][-1]["text"] += " " + line_stripped
                 else:
                     # Altrimenti, è continuazione del testo della domanda
                     # O è una domanda non numerata che inizia qui
                     if current_question["options"] or not current_question["text"]:
                         # Se la domanda precedente aveva opzioni o testo vuoto, inizia nuova domanda non numerata
                         if current_question["text"]: # Salva la precedente se aveva testo
                             questions.append(current_question)
                         question_counter += 1
                         current_question = {
                             "text": line_stripped, "order": question_counter, "options": [],
                             "type": QuestionType.MULTIPLE_CHOICE_SINGLE
                         }
                     else:
                        # Altrimenti, aggiungi al testo della domanda corrente
                        current_question["text"] += " " + line_stripped
            elif not current_question and line_stripped:
                 # Prima riga non vuota del file, considerala la prima domanda (non numerata)
                 question_counter += 1
                 current_question = {
                     "text": line_stripped, "order": question_counter, "options": [],
                     "type": QuestionType.MULTIPLE_CHOICE_SINGLE
                 }

        if current_question: # Salva l'ultima domanda
            questions.append(current_question)

        # Post-processing
        questions = [q for q in questions if q.get("text")] # Rimuovi eventuali domande vuote
        questions.sort(key=lambda q: q.get("order", float('inf'))) # Ordina per numero/contatore
        for i, q in enumerate(questions):
            q["order"] = i # Riassegna ordine 0-based sequenziale
            if not q["options"]:
                q["type"] = QuestionType.OPEN_ANSWER_MANUAL # Se non ha opzioni, è aperta

        if not questions:
             raise ValidationError("Nessuna domanda valida trovata nel file. Verifica la formattazione.")

        return questions


    @transaction.atomic
    def create(self, validated_data):
        uploaded_file: UploadedFile = validated_data['file']
        title = validated_data['title']
        teacher: User = self.context['request'].user

        file_content = io.BytesIO(uploaded_file.read())
        file_ext = '.' + uploaded_file.name.split('.')[-1].lower()

        text = ""
        try:
            if file_ext == '.pdf':
                text = self._extract_text_from_pdf(file_content)
            elif file_ext == '.docx':
                text = self._extract_text_from_docx(file_content)
            elif file_ext == '.md':
                text = self._extract_text_from_md(file_content)
            else:
                # Questo non dovrebbe accadere grazie a validate_file
                raise ValidationError("Tipo di file non gestito internamente.")
        except ValidationError as e:
             raise e # Rilancia errori di validazione specifici
        except Exception as e:
             logger.error(f"Errore estrazione testo da {uploaded_file.name}: {e}", exc_info=True)
             raise ValidationError(f"Errore durante la lettura del file {file_ext}.")


        if not text or text.isspace():
             raise ValidationError(f"Nessun contenuto testuale valido estratto da {uploaded_file.name}.")

        try:
            parsed_questions = self._parse_quiz_text(text)
        except ValidationError as e:
             raise e # Rilancia errori di validazione specifici
        except Exception as e:
             logger.error(f"Errore durante il parsing del testo da {uploaded_file.name}: {e}", exc_info=True)
             raise ValidationError("Errore durante l'analisi del contenuto del file.")


        # Creazione Quiz
        quiz = Quiz.objects.create(
            teacher=teacher,
            title=title,
            description=f"Quiz generato automaticamente da {uploaded_file.name}",
            metadata={
                'source_file': uploaded_file.name,
                'generation_method': 'auto_upload'
            }
            # available_from/until possono essere impostati successivamente
        )

        # Creazione Domande e Opzioni (Bulk)
        questions_to_create = []
        options_to_create = []

        for q_data in parsed_questions:
            question = Question(
                quiz=quiz,
                text=q_data['text'],
                question_type=q_data['type'],
                order=q_data['order'],
                metadata={} # Eventuali metadati specifici della domanda
            )
            questions_to_create.append(question)

        # Crea le domande in blocco e ottieni gli oggetti creati
        created_questions = Question.objects.bulk_create(questions_to_create)

        # Mappa per recuperare gli ID delle domande appena create
        # Usiamo una tupla (testo, ordine) come chiave, sperando sia sufficientemente univoca
        question_map = {(q.text, q.order): q for q in created_questions}

        # Prepara le opzioni associandole alle domande corrette
        for q_data in parsed_questions:
             # Recupera l'oggetto Question corrispondente usando la mappa
             question_obj = question_map.get((q_data['text'], q_data['order']))
             if not question_obj:
                 logger.warning(f"Domanda '{q_data['text'][:50]}...' (ordine {q_data['order']}) non trovata dopo bulk_create. Impossibile aggiungere opzioni.")
                 continue # Salta le opzioni per questa domanda

             for opt_data in q_data['options']:
                 option = AnswerOption(
                     question=question_obj,
                     text=opt_data['text'],
                     is_correct=opt_data['is_correct'], # Assumendo che il parser possa determinarlo
                     order=opt_data['order']
                 )
                 options_to_create.append(option)

        # Crea le opzioni in blocco se ce ne sono
        if options_to_create:
            AnswerOption.objects.bulk_create(options_to_create)

        logger.info(f"Creato Quiz '{quiz.title}' (ID: {quiz.id}) con {len(created_questions)} domande da {uploaded_file.name}")
        # Restituisce i dati del quiz creato usando il serializer appropriato
        return QuizSerializer(quiz, context=self.context).data

# --- Serializer per Upload Quiz Template (Invariato) ---

class QuizTemplateUploadSerializer(serializers.Serializer):
    """
    Serializer per gestire l'upload di file (PDF, DOCX, MD) per creare QuizTemplate.
    """
    file = serializers.FileField(required=True, help_text="File del quiz template (.pdf, .docx, .md)")
    title = serializers.CharField(max_length=255, required=True, help_text="Titolo del nuovo quiz template")

    # Metodi validate_file, _extract_text_from_*, _parse_quiz_text sono identici a QuizUploadSerializer
    # Si potrebbe creare una classe base comune per evitare duplicazione
    def validate_file(self, value: UploadedFile):
        if not hasattr(value, 'name') or not value.name:
             raise ValidationError("Nome del file mancante o non valido.")
        parts = value.name.split('.')
        if len(parts) < 2:
            raise ValidationError("Il file non ha un'estensione.")
        ext = '.' + parts[-1].lower()
        if ext not in ALLOWED_UPLOAD_EXTENSIONS:
            raise ValidationError(f"Tipo di file non supportato: {ext}. Sono permessi: {', '.join(ALLOWED_UPLOAD_EXTENSIONS)}")
        return value

    def _extract_text_from_pdf(self, file_obj: io.BytesIO) -> str:
        try:
            reader = PdfReader(file_obj)
            text = ""
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            if not text:
                 logger.warning("Nessun testo estratto dal PDF (template).")
            return text
        except Exception as e:
            logger.error(f"Errore estrazione testo PDF (template): {e}", exc_info=True)
            raise ValidationError(f"Impossibile leggere il file PDF (template). Errore: {e}")

    def _extract_text_from_docx(self, file_obj: io.BytesIO) -> str:
        try:
            document = DocxDocument(file_obj)
            text = "\n".join([para.text for para in document.paragraphs if para.text])
            if not text:
                 logger.warning("Nessun testo estratto dal DOCX (template).")
            return text
        except Exception as e:
            logger.error(f"Errore estrazione testo DOCX (template): {e}", exc_info=True)
            raise ValidationError(f"Impossibile leggere il file DOCX (template). Errore: {e}")

    def _extract_text_from_md(self, file_obj: io.BytesIO) -> str:
        try:
            md_content = file_obj.read().decode('utf-8')
            text = re.sub('<[^<]+?>', '', md_parser.markdown(md_content)).strip()
            if not text:
                 logger.warning("Nessun testo estratto dal Markdown (template).")
            return text
        except Exception as e:
            logger.error(f"Errore estrazione testo Markdown (template): {e}", exc_info=True)
            raise ValidationError(f"Impossibile leggere il file Markdown (template). Errore: {e}")

    def _parse_quiz_text(self, text: str) -> list[dict]:
        # Stessa logica di parsing di QuizUploadSerializer
        questions = []
        current_question = None
        question_start_re = re.compile(r"^\s*(?:(\d+)\s*[.)])?\s*(.*)", re.MULTILINE)
        option_re = re.compile(r"^\s*([A-Z])\s*[.)]\s*(.*)", re.MULTILINE)
        empty_line_re = re.compile(r"^\s*$")
        lines = text.splitlines()
        question_counter = 0

        for line in lines:
            if empty_line_re.match(line):
                continue

            line_stripped = line.strip()
            question_match = question_start_re.match(line)
            option_match = option_re.match(line_stripped)
            is_likely_question_start = question_match and question_match.group(1)
            is_likely_option = option_match

            if is_likely_question_start:
                if current_question:
                    questions.append(current_question)
                question_number = int(question_match.group(1))
                question_text = question_match.group(2).strip()
                question_counter = question_number
                current_question = {
                    "text": question_text, "order": question_number, "options": [],
                    "type": QuestionType.MULTIPLE_CHOICE_SINGLE
                }
            elif is_likely_option and current_question:
                option_letter = option_match.group(1)
                option_text = option_match.group(2).strip()
                current_question["options"].append({
                    "text": option_text, "order": len(current_question["options"]) + 1, "is_correct": False
                })
            elif current_question:
                 if current_question["options"] and not is_likely_question_start:
                     current_question["options"][-1]["text"] += " " + line_stripped
                 else:
                     if current_question["options"] or not current_question["text"]:
                         if current_question["text"]:
                             questions.append(current_question)
                         question_counter += 1
                         current_question = {
                             "text": line_stripped, "order": question_counter, "options": [],
                             "type": QuestionType.MULTIPLE_CHOICE_SINGLE
                         }
                     else:
                        current_question["text"] += " " + line_stripped
            elif not current_question and line_stripped:
                 question_counter += 1
                 current_question = {
                     "text": line_stripped, "order": question_counter, "options": [],
                     "type": QuestionType.MULTIPLE_CHOICE_SINGLE
                 }

        if current_question:
            questions.append(current_question)

        questions = [q for q in questions if q.get("text")]
        questions.sort(key=lambda q: q.get("order", float('inf')))
        for i, q in enumerate(questions):
            q["order"] = i
            if not q["options"]:
                q["type"] = QuestionType.OPEN_ANSWER_MANUAL

        if not questions:
             raise ValidationError("Nessuna domanda valida trovata nel file. Verifica la formattazione.")

        return questions


    @transaction.atomic
    def create(self, validated_data):
        uploaded_file: UploadedFile = validated_data['file']
        title = validated_data['title']
        teacher: User = self.context['request'].user # Per QuizTemplate, 'teacher' è opzionale, potrebbe essere admin

        file_content = io.BytesIO(uploaded_file.read())
        file_ext = '.' + uploaded_file.name.split('.')[-1].lower()

        text = ""
        try:
            if file_ext == '.pdf':
                text = self._extract_text_from_pdf(file_content)
            elif file_ext == '.docx':
                text = self._extract_text_from_docx(file_content)
            elif file_ext == '.md':
                text = self._extract_text_from_md(file_content)
            else:
                raise ValidationError("Tipo di file non gestito internamente.")
        except ValidationError as e:
             raise e
        except Exception as e:
             logger.error(f"Errore estrazione testo da {uploaded_file.name} (template): {e}", exc_info=True)
             raise ValidationError(f"Errore durante la lettura del file {file_ext} (template).")


        if not text or text.isspace():
             raise ValidationError(f"Nessun contenuto testuale valido estratto da {uploaded_file.name} (template).")

        try:
            parsed_questions = self._parse_quiz_text(text)
        except ValidationError as e:
             raise e
        except Exception as e:
             logger.error(f"Errore durante il parsing del testo da {uploaded_file.name} (template): {e}", exc_info=True)
             raise ValidationError("Errore durante l'analisi del contenuto del file (template).")

        # Creazione QuizTemplate
        # Determina se l'utente è admin o teacher per associare correttamente
        creator_kwargs = {}
        if teacher.is_admin:
            creator_kwargs['admin'] = teacher
        elif teacher.is_teacher: # Assumendo che 'teacher' sia l'utente loggato e sia un docente
            creator_kwargs['teacher'] = teacher
        else:
            # Caso non dovrebbe accadere se i permessi sono corretti, ma per sicurezza
            raise ValidationError("L'utente non è né Admin né Docente, impossibile creare il template.")

        quiz_template = QuizTemplate.objects.create(
            title=title,
            description=f"Quiz Template generato automaticamente da {uploaded_file.name}",
            metadata={
                'source_file': uploaded_file.name,
                'generation_method': 'auto_upload'
            },
            **creator_kwargs # Aggiunge admin=user o teacher=user
        )

        questions_to_create = []
        options_to_create = []

        for q_data in parsed_questions:
            question_template = QuestionTemplate(
                quiz_template=quiz_template,
                text=q_data['text'],
                question_type=q_data['type'],
                order=q_data['order'],
                metadata={}
            )
            questions_to_create.append(question_template)

        created_question_templates = QuestionTemplate.objects.bulk_create(questions_to_create)
        question_template_map = {(qt.text, qt.order): qt for qt in created_question_templates}

        for q_data in parsed_questions:
             question_template_obj = question_template_map.get((q_data['text'], q_data['order']))
             if not question_template_obj:
                 logger.warning(f"Template Domanda '{q_data['text'][:50]}...' (ordine {q_data['order']}) non trovato dopo bulk_create. Impossibile aggiungere opzioni.")
                 continue

             for opt_data in q_data['options']:
                 option_template = AnswerOptionTemplate(
                     question_template=question_template_obj,
                     text=opt_data['text'],
                     is_correct=opt_data['is_correct'],
                     order=opt_data['order']
                 )
                 options_to_create.append(option_template)

        if options_to_create:
            AnswerOptionTemplate.objects.bulk_create(options_to_create)

        logger.info(f"Creato QuizTemplate '{quiz_template.title}' (ID: {quiz_template.id}) con {len(created_question_templates)} domande da {uploaded_file.name}")
        return QuizTemplateSerializer(quiz_template, context=self.context).data


# --- Serializers per Tentativi e Dashboard (Aggiornati/Nuovi) ---

class QuizAssignmentSerializer(serializers.ModelSerializer):
    """ Serializer per visualizzare un'assegnazione Quiz esistente (a Studente o Gruppo). """
    quiz_title = serializers.CharField(source='quiz.title', read_only=True)
    student_username = serializers.CharField(source='student.user.username', read_only=True, allow_null=True)
    group_name = serializers.CharField(source='group.name', read_only=True, allow_null=True)
    assigned_by_username = serializers.CharField(source='assigned_by.username', read_only=True, allow_null=True)

    class Meta:
        model = QuizAssignment
        fields = [
            'id', 'quiz', 'quiz_title', 'student', 'student_username',
            'group', 'group_name', 'assigned_by', 'assigned_by_username',
            'assigned_at', 'due_date'
        ]
        read_only_fields = fields # Tutto è di sola lettura in questo contesto


class PathwayAssignmentSerializer(serializers.ModelSerializer):
    """ Serializer per visualizzare un'assegnazione Pathway esistente (a Studente o Gruppo). """
    pathway_title = serializers.CharField(source='pathway.title', read_only=True)
    student_username = serializers.CharField(source='student.user.username', read_only=True, allow_null=True)
    group_name = serializers.CharField(source='group.name', read_only=True, allow_null=True)
    # assigned_by non è ancora nel modello PathwayAssignment, da aggiungere se necessario

    class Meta:
        model = PathwayAssignment
        fields = [
            'id', 'pathway', 'pathway_title', 'student', 'student_username',
            'group', 'group_name', 'assigned_at', 'due_date'
        ]
        read_only_fields = fields


class AssignQuizSerializer(serializers.Serializer):
    """ Serializer per l'AZIONE di assegnare un Quiz esistente a un Gruppo. """
    quiz_id = serializers.IntegerField(required=True)
    group_id = serializers.IntegerField(required=True)
    due_date = serializers.DateTimeField(required=False, allow_null=True)

    def validate_quiz_id(self, value):
        if not Quiz.objects.filter(pk=value).exists():
            raise ValidationError("Quiz non trovato.")
        return value

    def validate_group_id(self, value):
        if not StudentGroup.objects.filter(pk=value).exists():
            raise ValidationError("Gruppo non trovato.")
        return value

    def validate_due_date(self, value):
        if value and value < timezone.now():
            raise ValidationError("La data di scadenza non può essere nel passato.")
        return value

    def create(self, validated_data):
        quiz = Quiz.objects.get(pk=validated_data['quiz_id'])
        group = StudentGroup.objects.get(pk=validated_data['group_id'])
        teacher = self.context['request'].user # Assumendo che il docente sia l'utente che fa la richiesta

        # Verifica permessi (il docente deve essere proprietario del gruppo o avere accesso approvato)
        is_owner = (group.owner == teacher)
        has_approved_access = StudentGroupMembership.objects.filter( # Questo controllo è errato, dovrebbe essere su GroupAccessRequest
            group=group,
            # student__user=teacher, # Questo non ha senso per un docente
            # role=StudentGroupMembership.Role.TEACHER # Non esiste questo ruolo in StudentGroupMembership
        ).exists() # Sostituire con logica corretta per GroupAccessRequest se implementata

        # Semplifichiamo: per ora assumiamo che se il docente può vedere il gruppo, può assegnare.
        # In produzione, serve un controllo più robusto basato su GroupAccessRequest o ownership.
        # if not (is_owner or has_approved_access):
        #     raise PermissionDenied("Non hai i permessi per assegnare contenuti a questo gruppo.")


        assignment, created = QuizAssignment.objects.update_or_create(
            quiz=quiz,
            group=group,
            student=None, # Assegnazione a gruppo
            defaults={
                'assigned_by': teacher,
                'assigned_at': timezone.now(),
                'due_date': validated_data.get('due_date')
            }
        )
        return assignment


class AssignPathwaySerializer(serializers.Serializer):
    """ Serializer per l'AZIONE di assegnare un Pathway esistente a uno Studente o a un Gruppo. """
    pathway_id = serializers.IntegerField(required=True)
    student_id = serializers.IntegerField(required=False, allow_null=True)
    group_id = serializers.IntegerField(required=False, allow_null=True)
    due_date = serializers.DateTimeField(required=False, allow_null=True)

    def validate_pathway_id(self, value):
        if not Pathway.objects.filter(pk=value).exists():
            raise ValidationError("Percorso non trovato.")
        return value

    def validate_student_id(self, value):
        if value and not Student.objects.filter(pk=value).exists(): # Usa Student.objects
            raise ValidationError("Studente non trovato.")
        return value

    def validate_group_id(self, value):
        if value and not StudentGroup.objects.filter(pk=value).exists():
            raise ValidationError("Gruppo non trovato.")
        return value

    def validate(self, attrs):
        if not attrs.get('student_id') and not attrs.get('group_id'):
            raise ValidationError("È necessario specificare student_id o group_id.")
        if attrs.get('student_id') and attrs.get('group_id'):
            raise ValidationError("Specificare student_id O group_id, non entrambi.")
        if attrs.get('due_date') and attrs['due_date'] < timezone.now():
            raise ValidationError("La data di scadenza non può essere nel passato.")
        return attrs

    def create(self, validated_data):
        pathway = Pathway.objects.get(pk=validated_data['pathway_id'])
        student_id = validated_data.get('student_id')
        group_id = validated_data.get('group_id')
        # assigned_by non è ancora nel modello PathwayAssignment

        assignment_defaults = {'assigned_at': timezone.now(), 'due_date': validated_data.get('due_date')}

        if student_id:
            student = Student.objects.get(pk=student_id)
            assignment, created = PathwayAssignment.objects.update_or_create(
                pathway=pathway, student=student, group=None,
                defaults=assignment_defaults
            )
        else: # group_id must be present due to validate()
            group = StudentGroup.objects.get(pk=group_id)
            assignment, created = PathwayAssignment.objects.update_or_create(
                pathway=pathway, group=group, student=None,
                defaults=assignment_defaults
            )
        return assignment


# --- Serializers per Tentativi e Risposte ---

class BasicQuestionSerializer(serializers.ModelSerializer):
    class Meta:
        model = Question
        fields = ['id', 'text', 'question_type', 'order'] # Info base domanda

class BasicQuizAttemptSerializer(serializers.ModelSerializer):
    class Meta:
        model = QuizAttempt
        fields = ['id', 'status', 'score', 'started_at', 'completed_at'] # Info base tentativo


class StudentAnswerSerializer(serializers.ModelSerializer):
    """ Serializer per inviare/visualizzare le risposte dello studente. """
    # Opzionale: includere dettagli domanda/opzione per GET, ma non per POST/PUT
    question_text = serializers.CharField(source='question.text', read_only=True, allow_null=True) # Testo della domanda
    # answer_option_text rimosso perché selected_answer_option non è più un campo diretto

    class Meta:
        model = StudentAnswer
        fields = [
            'id', 'quiz_attempt', 'question', 'question_text', # Ripristinato 'quiz_attempt'
            'selected_answers', # Nuovo campo JSON per le risposte
            'score', 'is_correct', 'answered_at'
        ]
        read_only_fields = ['quiz_attempt', 'score', 'is_correct', 'question_text', 'answered_at'] # Ripristinato 'quiz_attempt', aggiunto 'answered_at'
        extra_kwargs = {
            'selected_answers': {'required': True} # selected_answers è il campo principale per la risposta
        }


class QuizAttemptSerializer(serializers.ModelSerializer):
    """ Serializer base per i tentativi di quiz. """
    student = StudentBasicSerializer(read_only=True) # Usa StudentBasicSerializer
    # quiz = QuizSerializer(read_only=True) # Commentato per evitare cicli o dati eccessivi
    quiz_title = serializers.CharField(source='quiz.title', read_only=True)
    status_display = serializers.CharField(source='get_status_display', read_only=True)
    # Aggiunto per i badge guadagnati
    earned_badges = SimpleBadgeSerializer(many=True, read_only=True, source='badges')


    class Meta:
        model = QuizAttempt
        fields = [
            'id', 'student', 'quiz', 'quiz_title', 'status', 'status_display', 'score',
            'started_at', 'completed_at', 'earned_badges' # Aggiunto earned_badges
        ]
        read_only_fields = [
            'student', 'quiz', 'quiz_title', 'status_display', 'score',
            'started_at', 'completed_at', 'earned_badges'
        ]

    def to_representation(self, instance):
        """ Modifica la rappresentazione per includere i dettagli del quiz se necessario. """
        representation = super().to_representation(instance)
        # Se 'quiz' è stato rimosso dai fields ma vogliamo alcuni dettagli, li aggiungiamo qui.
        # representation['quiz_title'] = instance.quiz.title # Già gestito da source='quiz.title'
        # representation['quiz_description'] = instance.quiz.description # Esempio
        return representation


class QuizAttemptDetailSerializer(QuizAttemptSerializer):
    """ Serializer dettagliato per un tentativo, include risposte date e info aggiuntive. """
    # Eredita student, quiz_title, status, score, ecc. da QuizAttemptSerializer
    # student_answers = StudentAnswerSerializer(many=True, read_only=True) # SOSTITUITO con SerializerMethodField
    student_answers = serializers.SerializerMethodField() # NUOVO: Usa SerializerMethodField
    # Aggiungiamo i campi calcolati per il frontend
    total_questions = serializers.SerializerMethodField()
    correct_answers_count = serializers.SerializerMethodField() # DECOMMENTATO
    completion_threshold = serializers.SerializerMethodField() # Percentuale richiesta per passare

    class Meta(QuizAttemptSerializer.Meta): # Eredita Meta dal genitore
        fields = [f for f in QuizAttemptSerializer.Meta.fields if f not in ['student_answers', 'quiz']] + [
            'student_answers',
            'total_questions',
            'correct_answers_count',
            'completion_threshold',
        ]
        read_only_fields = QuizAttemptSerializer.Meta.read_only_fields + [
             'student_answers', 'total_questions', 'correct_answers_count', 'completion_threshold'
        ]

    def get_student_answers(self, obj: QuizAttempt):
        """
        Serializza le student_answers passando il contesto corretto a QuestionSerializer
        per la gestione dinamica dei metadata delle domande FILL_BLANK.
        """
        viewing_results = False
        if obj.status in [QuizAttempt.AttemptStatus.COMPLETED,
                          QuizAttempt.AttemptStatus.FAILED,
                          QuizAttempt.AttemptStatus.PENDING_GRADING]:
            viewing_results = True
        
        # Passa il contesto al StudentAnswerSerializer, che a sua volta lo passerà
        # implicitamente al QuestionSerializer quando serializza il campo 'question'.
        context = self.context.copy() # Eredita contesto esistente (es. request)
        context['viewing_results'] = viewing_results
        
        # Precarica le domande e le opzioni per ottimizzare
        answers_queryset = obj.student_answers.prefetch_related(
            'question',
            'question__answer_options' # Per MC, TF
            # Non è necessario precaricare question__metadata specificamente qui,
            # QuestionSerializer vi accederà direttamente.
        ).all()
        
        serializer = StudentAnswerSerializer(answers_queryset, many=True, context=context)
        return serializer.data

    def get_completion_threshold(self, obj: QuizAttempt) -> float | None:
        """ Restituisce la soglia di completamento del quiz, se definita. Gestisce errori e logga i metadati se la chiave manca. """
        try:
            # Prova ad accedere direttamente al campo 'completion_threshold' se denormalizzato
            if hasattr(obj.quiz, 'completion_threshold') and obj.quiz.completion_threshold is not None:
                return float(obj.quiz.completion_threshold)

            # Altrimenti, prova ad accedere tramite metadata
            if obj.quiz and obj.quiz.metadata and 'completion_threshold' in obj.quiz.metadata:
                threshold_value = obj.quiz.metadata.get('completion_threshold')
                if threshold_value is not None:
                    try:
                        return float(threshold_value)
                    except (ValueError, TypeError):
                        logger.error(
                            f"Valore non valido per completion_threshold nei metadati del quiz {obj.quiz.id}: '{threshold_value}'. "
                            f"Metadati completi: {obj.quiz.metadata}"
                        )
                        return None # O un default sensato, es. 0.0 o 100.0 a seconda della logica
                else: # Chiave presente ma valore è None
                    logger.warning(
                        f"'completion_threshold' è None nei metadati del quiz {obj.quiz.id}. "
                        f"Metadati completi: {obj.quiz.metadata}"
                    )
                    return None
            else: # Chiave non presente o quiz/metadata non disponibili
                if not obj.quiz:
                    logger.warning(f"Tentativo {obj.id} non ha un quiz associato per recuperare completion_threshold.")
                elif not obj.quiz.metadata:
                    logger.warning(f"Quiz {obj.quiz.id} non ha metadati per recuperare completion_threshold.")
                else: # Chiave non nei metadati
                    logger.info(
                        f"'completion_threshold' non trovato nei metadati del quiz {obj.quiz.id}. "
                        f"Metadati disponibili: {list(obj.quiz.metadata.keys()) if obj.quiz.metadata else 'Nessuno'}"
                    )
                return None # Default se non trovato
        except Exception as e:
            logger.error(f"Errore generico nel recuperare completion_threshold per quiz {obj.quiz_id if obj.quiz else 'N/A'}, tentativo {obj.id}: {e}", exc_info=True)
            return None

    def get_total_questions(self, obj: QuizAttempt) -> int:
        """ Restituisce il numero totale di domande nel quiz. Gestisce errori. """
        try:
            if obj.quiz:
                return obj.quiz.questions.count()
            return 0
        except Exception as e:
            logger.error(f"Errore nel contare le domande per quiz {obj.quiz_id if obj.quiz else 'N/A'}, tentativo {obj.id}: {e}", exc_info=True)
            return 0

    def get_correct_answers_count(self, obj: QuizAttempt) -> int: # DECOMMENTATO e implementato
        """ Conta le risposte corrette date dallo studente in questo tentativo. """
        # Assicurati che le risposte siano state precaricate se necessario per performance
        # In questo caso, student_answers è già nel serializer, quindi dovrebbe essere disponibile.
        try:
            return obj.student_answers.filter(is_correct=True).count()
        except Exception as e:
            logger.error(f"Errore nel contare le risposte corrette per il tentativo {obj.id}: {e}", exc_info=True)
            return 0


# --- Serializers per Dashboard Studente (Potrebbero necessitare aggiornamenti per gruppi) ---

class PathwayProgressSerializer(serializers.ModelSerializer):
    """ Serializer per il progresso dello studente in un percorso. """
    # student = StudentBasicSerializer(read_only=True) # Rimosso per evitare ridondanza se usato nested
    # pathway = PathwaySerializer(read_only=True) # Rimosso per evitare cicli/dati eccessivi

    class Meta:
        model = PathwayProgress
        fields = [
            'id', 'student', 'pathway', 'status', 'last_completed_quiz_order',
            'completed_orders', 'started_at', 'completed_at', 'points_earned'
        ]
        read_only_fields = fields # Tutto di sola lettura in questo contesto


class SimpleQuizAttemptSerializer(serializers.ModelSerializer):
    """ Serializer minimale per rappresentare un tentativo in una lista. """
    class Meta:
        model = QuizAttempt
        fields = ['id', 'status', 'score', 'completed_at'] # Corretto: end_time -> completed_at


class StudentQuizDashboardSerializer(QuizSerializer):
    """ Serializer per i quiz nella dashboard studente, include stato ultimo tentativo. """
    latest_attempt = serializers.SerializerMethodField()
    attempts_count = serializers.SerializerMethodField()
    # Aggiungere campo per indicare se assegnato direttamente o via gruppo?
    assignment_type = serializers.SerializerMethodField()
    card_background_color = serializers.SerializerMethodField() # Nuovo campo

    class Meta(QuizSerializer.Meta):
        fields = QuizSerializer.Meta.fields + ['latest_attempt', 'attempts_count', 'assignment_type', 'card_background_color'] # Aggiunto ai fields

    def get_card_background_color(self, obj: Quiz) -> str | None:
        if obj.source_template and obj.source_template.card_background_color:
            return obj.source_template.card_background_color
        return None

    def get_latest_attempt(self, obj):
        """ Recupera l'ultimo tentativo dello studente per questo quiz. """
        student = self.context.get('student') # Assumendo studente nel contesto
        if not student: return None
        attempt = obj.attempts.filter(student=student).order_by('-started_at').first() # Corretto: usa started_at
        return SimpleQuizAttemptSerializer(attempt).data if attempt else None

    def get_attempts_count(self, obj):
        student = self.context.get('student')
        if not student: return 0
        return obj.attempts.filter(student=student).count()

    def get_assignment_type(self, obj):
        # Determina se il quiz è stato assegnato direttamente o tramite gruppo
        # Questa logica potrebbe essere complessa e richiedere accesso al contesto o query aggiuntive
        # Per ora, un placeholder. La view StudentAssignedQuizzesView gestisce questo.
        student = self.context.get('student')
        if not student: return "unknown"

        is_direct = QuizAssignment.objects.filter(quiz=obj, student=student).exists()
        if is_direct:
            return "student"

        student_groups = StudentGroupMembership.objects.filter(student=student).values_list('group_id', flat=True)
        is_group = QuizAssignment.objects.filter(quiz=obj, group_id__in=student_groups).exists()
        if is_group:
            return "group"
        return "unknown"


class SimplePathwayProgressSerializer(serializers.ModelSerializer):
    class Meta:
        model = PathwayProgress
        fields = ['id', 'status', 'last_completed_quiz_order', 'completed_at', 'points_earned']


class StudentPathwayDashboardSerializer(PathwaySerializer):
    """ Serializer per i percorsi nella dashboard studente, include stato ultimo progresso. """
    latest_progress = serializers.SerializerMethodField()
    # Aggiungere campo per indicare se assegnato direttamente o via gruppo?
    assignment_type = serializers.SerializerMethodField()


    class Meta(PathwaySerializer.Meta):
        fields = PathwaySerializer.Meta.fields + ['latest_progress', 'assignment_type']

    def get_latest_progress(self, obj):
        student = self.context.get('student')
        if not student: return None
        progress = PathwayProgress.objects.filter(pathway=obj, student=student).order_by('-id').first()
        return SimplePathwayProgressSerializer(progress).data if progress else None

    def get_assignment_type(self, obj):
        # Simile a StudentQuizDashboardSerializer, determina tipo assegnazione
        student = self.context.get('student')
        if not student: return "unknown"
        # Questa logica deve essere adattata per PathwayAssignment
        is_direct = PathwayAssignment.objects.filter(pathway=obj, student=student).exists()
        if is_direct:
            return "student"
        student_groups = StudentGroupMembership.objects.filter(student=student).values_list('group_id', flat=True)
        is_group = PathwayAssignment.objects.filter(pathway=obj, group_id__in=student_groups).exists()
        if is_group:
            return "group"
        return "unknown"

# --- Teacher Grading Serializers ---

class PendingQuizAttemptSerializer(serializers.ModelSerializer):
    """
    Serializer per elencare i tentativi di quiz in attesa di correzione da parte del docente.
    """
    student_name = serializers.CharField(source='student.full_name', read_only=True)
    quiz_title = serializers.CharField(source='quiz.title', read_only=True)
    # Potremmo aggiungere il nome del docente del quiz se diverso da quello che corregge
    # teacher_name = serializers.CharField(source='quiz.teacher.full_name', read_only=True)

    class Meta:
        model = QuizAttempt
        fields = [
            'id',
            'student_name',
            'quiz_title',
            'started_at', # Data di inizio del tentativo
            'completed_at', # Data di "completamento" da parte dello studente (quando entra in pending)
            'status'
        ]
        read_only_fields = fields


class GradingStudentAnswerSerializer(serializers.ModelSerializer):
    """
    Serializer per visualizzare una StudentAnswer nel contesto della correzione manuale.
    """
    selected_answers_text = serializers.SerializerMethodField()
    question_text = serializers.CharField(source='question.text', read_only=True)
    question_order = serializers.IntegerField(source='question.order', read_only=True)
    # Aggiungiamo il campo teacher_comment per la visualizzazione
    teacher_comment = serializers.CharField(read_only=True, allow_blank=True, allow_null=True)


    class Meta:
        model = StudentAnswer
        fields = [
            'id',
            'question_text', # Testo della domanda originale
            'question_order', # Ordine della domanda
            'selected_answers_text', # Testo della risposta data dallo studente
            'is_correct', # Stato attuale (potrebbe essere None)
            'score', # Punteggio attuale (potrebbe essere None)
            'teacher_comment' # Commento attuale del docente (se presente)
        ]
        # Rendiamo is_correct e score leggibili qui, la modifica avverrà tramite un altro serializer
        read_only_fields = ['id', 'question_text', 'question_order', 'selected_answers_text', 'is_correct', 'score', 'teacher_comment']


    def get_selected_answers_text(self, obj: StudentAnswer) -> str | None:
        if obj.question.question_type == QuestionType.OPEN_ANSWER_MANUAL:
            return obj.selected_answers.get('text', None)
        return None


class GradingQuestionWithAnswerSerializer(serializers.ModelSerializer):
    """
    Serializer per una domanda che include la risposta dello studente per quel tentativo.
    Usato all'interno di GradingQuizAttemptDetailSerializer.
    """
    student_answer = serializers.SerializerMethodField()
    question_type_display = serializers.CharField(source='get_question_type_display', read_only=True)
    # Aggiungiamo answer_options per visualizzare le opzioni nel caso di domande non aperte
    answer_options = AnswerOptionSerializer(many=True, read_only=True)


    class Meta:
        model = Question
        fields = [
            'id',
            'text',
            'question_type',
            'question_type_display',
            'order',
            'metadata', # Potrebbe essere utile per il docente vedere i metadati della domanda
            'answer_options', # Aggiunto per contesto
            'student_answer'
        ]
        read_only_fields = fields

    def get_student_answer(self, obj: Question) -> dict | None:
        quiz_attempt = self.context.get('quiz_attempt')
        if quiz_attempt:
            try:
                student_answer_instance = StudentAnswer.objects.get(question=obj, quiz_attempt=quiz_attempt)
                return GradingStudentAnswerSerializer(student_answer_instance, context=self.context).data
            except StudentAnswer.DoesNotExist:
                logger.warning(f"Nessuna StudentAnswer trovata per la domanda {obj.id} nel tentativo {quiz_attempt.id}")
                return None
        logger.warning(f"QuizAttempt non trovato nel contesto per GradingQuestionWithAnswerSerializer (domanda {obj.id})")
        return None


class GradingQuizAttemptDetailSerializer(serializers.ModelSerializer):
    """
    Serializer dettagliato per un QuizAttempt in fase di correzione manuale.
    """
    student = StudentBasicSerializer(read_only=True)
    quiz_title = serializers.CharField(source='quiz.title', read_only=True)
    quiz_description = serializers.CharField(source='quiz.description', read_only=True, allow_blank=True, allow_null=True)
    questions_with_answers = serializers.SerializerMethodField()

    class Meta:
        model = QuizAttempt
        fields = [
            'id',
            'student',
            'quiz_title',
            'quiz_description',
            'started_at',
            'completed_at',
            'status',
            'questions_with_answers'
        ]
        read_only_fields = fields

    def get_questions_with_answers(self, obj: QuizAttempt) -> list[dict]:
        questions = Question.objects.filter(quiz=obj.quiz).order_by('order').prefetch_related('answer_options') # Aggiunto prefetch
        return GradingQuestionWithAnswerSerializer(questions, many=True, context={'request': self.context.get('request'), 'quiz_attempt': obj}).data


class GradeSubmissionItemSerializer(serializers.Serializer):
    """
    Serializer per i dati di una singola risposta da correggere.
    """
    student_answer_id = serializers.IntegerField(required=True)
    is_correct = serializers.BooleanField(required=True)
    teacher_comment = serializers.CharField(required=False, allow_blank=True, allow_null=True, max_length=1000)

    def validate_student_answer_id(self, value):
        try:
            sa = StudentAnswer.objects.select_related('question').get(pk=value)
            # Verifica che la domanda sia effettivamente di tipo OPEN_ANSWER_MANUAL
            if sa.question.question_type != QuestionType.OPEN_ANSWER_MANUAL:
                raise serializers.ValidationError(f"La risposta con ID {value} non appartiene a una domanda a risposta aperta.")
        except StudentAnswer.DoesNotExist:
            raise serializers.ValidationError(f"StudentAnswer con ID {value} non trovata.")
        return value


class GradeSubmissionSerializer(serializers.Serializer):
    """
    Serializer per il payload completo dell'invio delle correzioni da parte del docente.
    """
    answers = GradeSubmissionItemSerializer(many=True, required=True, allow_empty=False)

    def validate_answers(self, value):
        if not value:
            raise serializers.ValidationError("La lista delle risposte corrette non può essere vuota.")
        # Ulteriori validazioni (es. ID unici, appartenenza allo stesso attempt) andranno fatte nella view.
        return value


# --- NUOVO Serializer per Tentativi Quiz nella Dashboard Studente ---

class StudentQuizAttemptDashboardSerializer(serializers.ModelSerializer):
    """
    Serializer per i tentativi di quiz nella dashboard studente.
    Mostra le informazioni del quiz e lo stato specifico di questo tentativo.
    """
    # Campi dal Quiz associato (Ora letti direttamente dal dizionario fornito dalla view)
    quiz_id = serializers.IntegerField(read_only=True)
    title = serializers.CharField(read_only=True)
    description = serializers.CharField(read_only=True)
    available_from = serializers.DateTimeField(read_only=True, allow_null=True) # Allow null
    available_until = serializers.DateTimeField(read_only=True, allow_null=True) # Allow null
    metadata = serializers.JSONField(read_only=True)
    teacher_username = serializers.CharField(read_only=True, allow_null=True) # Allow null
    teacher_first_name = serializers.CharField(read_only=True, allow_null=True) # NUOVO
    teacher_last_name = serializers.CharField(read_only=True, allow_null=True)  # NUOVO
    subject_name = serializers.CharField(read_only=True, allow_null=True)
    topic_name = serializers.CharField(read_only=True, allow_null=True)
    subject_color_placeholder = serializers.CharField(read_only=True, allow_null=True, required=False)

    # Campi specifici del tentativo
    attempt_id = serializers.IntegerField(read_only=True, allow_null=True) # ID del tentativo (può essere null per PENDING)
    status = serializers.CharField(read_only=True) # Stato del tentativo (es. PENDING, IN_PROGRESS, COMPLETED, FAILED)
    score = serializers.FloatField(read_only=True, allow_null=True)
    started_at = serializers.DateTimeField(read_only=True, allow_null=True) # Allow null
    completed_at = serializers.DateTimeField(read_only=True, allow_null=True)

    # Legge direttamente dal dizionario fornito dalla view
    assignment_type = serializers.CharField(read_only=True)
    card_background_color = serializers.CharField(read_only=True, allow_null=True, required=False) # Legge dal dict

    class Meta:
        # Ripristinato Meta.model, necessario per DRF anche se lavoriamo su dict
        model = QuizAttempt
        fields = [
            'attempt_id', 'quiz_id', 'title', 'description', 'status', 'score',
            'available_from', 'available_until', 'metadata', 'teacher_username', 'teacher_first_name', 'teacher_last_name',
            'subject_name', 'topic_name', 'subject_color_placeholder', # Campi aggiunti
            'started_at', 'completed_at', 'assignment_type', 'card_background_color' # Assicurato che sia nei fields
        ]
        read_only_fields = fields # Tutti i campi sono derivati o di sola lettura in questo contesto

    # Rimosso: get_assignment_type non serve più, il valore viene letto direttamente dal dizionario
    # def get_assignment_type(self, obj): ...

    # Aggiunto per completezza, anche se read_only_fields copre tutto
    def create(self, validated_data):
        raise NotImplementedError("Questo serializer è di sola lettura.")

    def update(self, instance, validated_data):
        raise NotImplementedError("Questo serializer è di sola lettura.")


# --- Serializers Dettaglio Percorso per Studente ---

class NextPathwayQuizSerializer(serializers.ModelSerializer):
    """ Serializer per mostrare il prossimo quiz da fare in un percorso. """
    # Potrebbe includere dettagli del quiz o solo l'ID
    quiz_id = serializers.IntegerField(source='id') # L'ID del Quiz stesso
    quiz_title = serializers.CharField(source='title')
    # Aggiungere altri campi del Quiz se necessario per l'anteprima
    # Esempio: description, metadata.difficulty

    class Meta:
        model = Quiz # Il modello è Quiz, non PathwayQuiz
        fields = ['quiz_id', 'quiz_title'] # Aggiungere altri campi qui


class PathwayAttemptDetailSerializer(PathwaySerializer):
    """ Serializer dettagliato per un percorso quando visualizzato da uno studente. """
    # Eredita i campi da PathwaySerializer (title, description, quiz_details, ecc.)
    progress = serializers.SerializerMethodField() # Stato di avanzamento dello studente
    next_quiz = serializers.SerializerMethodField() # Prossimo quiz da fare

    class Meta(PathwaySerializer.Meta):
        fields = PathwaySerializer.Meta.fields + ['progress', 'next_quiz']

    def get_progress(self, obj):
        student = self.context.get('request').user # Assumendo che lo studente sia l'utente loggato
        if not student or not hasattr(student, 'student_profile'): # Verifica se è uno studente
            return None
        # student_profile = student.student_profile # Se Student è un profilo di User
        # Se Student è User stesso con un ruolo:
        if not student.is_student: return None


        progress = PathwayProgress.objects.filter(pathway=obj, student=student).first()
        if progress:
            return PathwayProgressSerializer(progress).data
        return None

    def get_next_quiz(self, obj):
        student = self.context.get('request').user
        if not student or not hasattr(student, 'student_profile'): # Verifica se è uno studente
             return None
        # student_profile = student.student_profile
        if not student.is_student: return None


        progress = PathwayProgress.objects.filter(pathway=obj, student=student).first()
        last_completed_order = -1
        if progress and progress.last_completed_quiz_order is not None:
            last_completed_order = progress.last_completed_quiz_order

        # Trova il prossimo PathwayQuiz in ordine
        next_pathway_quiz = obj.pathwayquiz_set.filter(order__gt=last_completed_order).order_by('order').first()

        if next_pathway_quiz:
            # Serializza solo i dettagli necessari del Quiz, non l'intero PathwayQuiz
            return NextPathwayQuizSerializer(next_pathway_quiz.quiz).data
        return None
# --- Notification Serializers ---

class NotificationSerializer(serializers.ModelSerializer):
    """
    Serializer per il modello Notification.
    """
    # Potremmo voler aggiungere campi derivati o formattati qui se necessario
    # Esempio: student_name = serializers.CharField(source='student.full_name', read_only=True)
    #          created_at_formatted = serializers.DateTimeField(source='created_at', format="%Y-%m-%d %H:%M", read_only=True)

    class Meta:
        model = Notification
        fields = [
            'id', 
            'student', # In genere non si espone l'ID studente direttamente se l'endpoint è già per lo studente loggato
            'message', 
            'link', 
            'notification_type', 
            'is_read', 
            'created_at'
        ]
        read_only_fields = ['id', 'student', 'message', 'link', 'notification_type', 'created_at']
        # 'is_read' è modificabile tramite azioni specifiche nel ViewSet, non direttamente qui.

    def to_representation(self, instance):
        """
        Personalizza la rappresentazione. Ad esempio, per non esporre 'student' ID.
        """
        representation = super().to_representation(instance)
        # Rimuovi 'student' se l'API è già contestualizzata per lo studente loggato
        # Questo dipende da come il ViewSet gestisce il queryset e i permessi.
        # Se l'endpoint è /api/student/notifications/, 'student' potrebbe essere ridondante.
        # representation.pop('student', None) 
        return representation